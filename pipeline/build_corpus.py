#!/usr/bin/env python3
"""Corpus builder template — adapt to your data sources.

The orchestrator invokes this as `python build_corpus.py --name <corpus> --out <path>`.
You return a JSONL file with one record per line:

    {"text": "<chunk of text>", "source": "<source.tag>", "id": "<sha-prefix>"}

This file is a *template*. The BUILD_PLAN at the bottom and the source-loader
functions are the parts you customize for your data. The utilities at the top
(`chunk`, `text_id`, `write_jsonl`, `load_jsonl`) work for any project.

Default strategy: corpora grow over versions (v6 → v7 → v8 → ...). After the
final growth step, the orchestrator just iterates epochs on the largest corpus.
See README.md "Corpus growth then penetration" for the rationale.
"""
import argparse
import hashlib
import json
import os
import re
import sys
from pathlib import Path

# ============================================================
# Paths — adapt for your install layout
# ============================================================

ROOT = Path(os.environ.get("VOICE_LORA_ROOT", Path.home() / "voice-lora"))
RAW = ROOT / "raw"        # put your raw source data here (gitignored)
DATA = ROOT / "data"      # output corpora live here

# Chunking sizes — tuned for ~4K seq_len training. If your seq_len is larger,
# raise MAX_CHARS proportionally (Japanese ≈ 2 chars/token, English ≈ 4).
MIN_CHARS = 400
MAX_CHARS = 16000

# ============================================================
# Generic utilities — reusable across projects
# ============================================================

def load_jsonl(p):
    recs = []
    if not Path(p).exists():
        return recs
    with open(p, encoding="utf-8", errors="ignore") as f:
        for ln in f:
            ln = ln.strip()
            if not ln:
                continue
            try:
                recs.append(json.loads(ln))
            except json.JSONDecodeError:
                continue
    return recs


def chunk(text, max_chars=MAX_CHARS, min_chars=MIN_CHARS):
    """Split text on paragraph boundaries. Drop fragments < min_chars."""
    text = text.strip()
    if len(text) <= max_chars:
        return [text] if len(text) >= min_chars else []
    out, cur = [], ""
    for para in re.split(r"\n\n+", text):
        if len(cur) + len(para) + 2 > max_chars:
            if len(cur) >= min_chars:
                out.append(cur.strip())
            cur = para
        else:
            cur = cur + "\n\n" + para if cur else para
    if len(cur) >= min_chars:
        out.append(cur.strip())
    return out


def text_id(t):
    """Stable short id from first 512 chars. Used for dedup."""
    return hashlib.sha256(t[:512].encode("utf-8", errors="ignore")).hexdigest()[:16]


def write_jsonl(path, records):
    """Write records to JSONL with dedup-by-prefix-hash and min-length filtering."""
    n, total_chars = 0, 0
    seen = set()
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        for r in records:
            t = r.get("text", "").strip()
            if len(t) < MIN_CHARS:
                continue
            h = text_id(t)
            if h in seen:
                continue
            seen.add(h)
            r.setdefault("id", h)
            f.write(json.dumps(r, ensure_ascii=False) + "\n")
            n += 1
            total_chars += len(t)
    print(f"  wrote {path.name}: {n:,} records, {total_chars:,} chars "
          f"(~{total_chars/2_000_000:.1f}M tok JA, ~{total_chars/4_000_000:.1f}M tok EN)")
    return n, total_chars


# ============================================================
# Source loaders — CUSTOMIZE THESE FOR YOUR DATA
# ============================================================
#
# Each loader appends dicts to the `records` list with shape:
#   {"text": <str>, "source": "<Category.subtag>"}
# The source tag is purely informational (lets you grep / filter later).
#
# Below are example loaders that match common data shapes. Delete what you
# don't need, copy/adapt what you do.

def add_existing_jsonl(records, jsonl_path, source_tag, text_key="text"):
    """Read a pre-existing JSONL where each record already has a text field."""
    n = 0
    for r in load_jsonl(jsonl_path):
        t = (r.get(text_key) or "").strip()
        for c in chunk(t):
            records.append({"text": c, "source": source_tag})
            n += 1
    print(f"  + {source_tag} ({Path(jsonl_path).name}): {n:,} chunks")


def add_authored_jsonl(records, jsonl_path, source_tag,
                      body_key="body", author_predicate=None):
    """Read JSONL where some records are by the target author. `author_predicate`
    receives the raw record and returns True iff it should be included.

    Example use: Apple Mail dump (.jsonl from `mbox_to_jsonl`), Slack export,
    forum dump, etc.
    """
    n = 0
    for r in load_jsonl(jsonl_path):
        if author_predicate and not author_predicate(r):
            continue
        body = (r.get(body_key) or "").strip()
        # Strip quoted reply lines that start with ">"
        body = re.sub(r"^>.*$", "", body, flags=re.M)
        body = re.sub(r"\n{3,}", "\n\n", body).strip()
        for c in chunk(body):
            records.append({"text": c, "source": source_tag})
            n += 1
    print(f"  + {source_tag} ({Path(jsonl_path).name}): {n:,} chunks")


def add_text_glob(records, root_dir, glob_pattern, source_tag, paragraph_split=True):
    """Recursive glob over a directory, reading .txt/.md/etc into chunks.

    If paragraph_split=True, treats consecutive double-newlines as paragraph
    breaks (good for prose). If False, inserts paragraph breaks after sentence-
    end punctuation (good for transcripts/monologues with no formatting).
    """
    n = 0
    root_dir = Path(root_dir)
    if not root_dir.exists():
        print(f"  ! {source_tag}: {root_dir} missing, skipped", file=sys.stderr)
        return
    for p in root_dir.rglob(glob_pattern):
        try:
            t = p.read_text(encoding="utf-8", errors="ignore")
            if not paragraph_split:
                t = re.sub(r"([。．！？.!?])", r"\1\n", t)
                t = re.sub(r"\n{2,}", "\n\n", t)
            for c in chunk(t):
                records.append({"text": c, "source": source_tag})
                n += 1
        except Exception as e:
            print(f"  ! {p.name}: {e}", file=sys.stderr)
    print(f"  + {source_tag} ({root_dir.name}/{glob_pattern}): {n:,} chunks")


def add_docx_dir(records, root_dir, source_tag):
    """Extract text from all .docx in a directory."""
    n = 0
    root_dir = Path(root_dir)
    if not root_dir.exists():
        return
    try:
        from docx import Document
    except ImportError:
        print(f"  ! {source_tag}: python-docx missing, skipped", file=sys.stderr)
        return
    for docx in root_dir.rglob("*.docx"):
        try:
            doc = Document(str(docx))
            parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            for c in chunk("\n\n".join(parts)):
                records.append({"text": c, "source": source_tag})
                n += 1
        except Exception as e:
            print(f"  ! docx {docx.name}: {e}", file=sys.stderr)
    print(f"  + {source_tag}: {n:,} chunks")


def add_slack_export(records, jsonl_path, source_tag, author_keyword):
    """Slack JSONL where each message has user/text fields. Keeps only messages
    by users whose name/email contains `author_keyword` (case-insensitive).
    Cleans Slack markup tokens (<@U…>, <#C…|name>, <https://…|label>).
    """
    n = 0
    kw = author_keyword.lower()
    for r in load_jsonl(jsonl_path):
        user_str = (r.get("user") or r.get("user_name") or "").lower()
        if kw not in user_str:
            continue
        t = (r.get("text") or r.get("body") or "").strip()
        t = re.sub(r"<@[A-Z0-9]+>", "", t)
        t = re.sub(r"<#[A-Z0-9]+\|([^>]+)>", r"#\1", t)
        t = re.sub(r"<(https?://[^|>]+)\|([^>]+)>", r"\2", t)
        t = re.sub(r"<(https?://[^>]+)>", r"\1", t)
        if len(t) < 80:
            continue
        for c in chunk(t):
            records.append({"text": c, "source": source_tag})
            n += 1
    print(f"  + {source_tag} ({Path(jsonl_path).name}): {n:,} chunks")


# ============================================================
# BUILD PLAN — CUSTOMIZE THIS
# ============================================================
#
# Map corpus name → list of (loader_fn, args) to call in order.
# Names referenced in versions.yaml MUST match keys here.
#
# Convention: each version step adds one more bundle of sources on top of the
# previous. The orchestrator's `resume_from: previous` chains adapters; this
# function chains *corpora*.

def build_v6(records):
    """Initial seed corpus — usually a JSONL you've curated by hand or from
    your own preprocessing pipeline."""
    add_existing_jsonl(records, RAW / "seed_corpus.jsonl", source_tag="Seed")


def build_v7(records):
    """v6 + personal-channel additions (emails, chat logs, etc)."""
    build_v6(records)
    # Email: mbox dumped to JSONL with is_author_sent flag
    add_authored_jsonl(records, RAW / "email.jsonl",
                       source_tag="Personal.email",
                       body_key="body",
                       author_predicate=lambda r: r.get("is_author_sent"))
    # Slack: filter by author name keyword
    add_slack_export(records, RAW / "slack_team_a.jsonl",
                     source_tag="Personal.slack",
                     author_keyword="your_username")


def build_v8(records):
    """v7 + interview transcripts + presentation notes + docx archives."""
    build_v7(records)
    # Long-form interview transcripts (no paragraph formatting, sentence-split)
    add_text_glob(records, RAW / "interviews", "*.txt",
                  source_tag="Interview.transcripts",
                  paragraph_split=False)
    # Talk notes / scripts (markdown + plain text)
    add_text_glob(records, RAW / "presentations", "*.md",
                  source_tag="Presentation.notes")
    add_text_glob(records, RAW / "presentations", "*.txt",
                  source_tag="Presentation.notes")
    # Docx archive (book drafts, manuscripts, etc)
    add_docx_dir(records, RAW / "manuscripts", source_tag="Manuscript.docx")


BUILD_PLAN = {
    "v6": build_v6,
    "v7": build_v7,
    "v8": build_v8,
    # add more entries for v9, v10... if you want corpora to keep growing.
    # The orchestrator only calls this builder when a corpus_v*.jsonl doesn't
    # exist yet — so adding entries here lets you grow further without
    # invalidating earlier corpora.
}


# ============================================================
# CLI
# ============================================================

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--name", required=True, help="corpus name (e.g. v7, v8)")
    ap.add_argument("--out", required=True, help="output JSONL path")
    args = ap.parse_args()

    if args.name not in BUILD_PLAN:
        print(f"ERROR: corpus '{args.name}' not in BUILD_PLAN. "
              f"Known: {sorted(BUILD_PLAN)}", file=sys.stderr)
        sys.exit(1)

    print(f"=== Building corpus {args.name} -> {args.out} ===")
    records = []
    BUILD_PLAN[args.name](records)
    write_jsonl(args.out, records)


if __name__ == "__main__":
    main()
